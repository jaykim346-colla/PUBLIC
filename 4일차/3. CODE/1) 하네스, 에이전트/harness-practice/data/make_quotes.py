# -*- coding: utf-8 -*-
"""
견적 자동 생성 프로그램
아래 4단계를 항상 이 순서 그대로 실행한다. (판단/순서변경 없음)
  1. 요청서에서 출발·도착·타입·수량 읽기
  2. 운임표에서 같은 구간+타입 단가 찾기
  3. 합계 = 단가 × 수량 으로 계산
  4. 양식 빈칸 채워 건별 견적서로 저장
"""

import os
from datetime import date
import openpyxl
from docx import Document

# ── 파일 경로 (스크립트와 같은 폴더 기준) ──────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
SHIPPER_REQUEST = os.path.join(BASE, "shipper_request.xlsx")
RATE_TABLE      = os.path.join(BASE, "rate_table.xlsx")
QUOTE_TEMPLATE  = os.path.join(BASE, "quote_template.docx")
OUTPUT_DIR      = os.path.join(BASE, "quotes_out")


# ── 단계 1: 요청서에서 출발·도착·타입·수량 읽기 ────────────────────
def step1_read_requests(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    header = rows[0]
    idx = {name: i for i, name in enumerate(header)}
    requests = []
    for r in rows[1:]:
        if r[idx["요청번호"]] is None:
            continue
        requests.append({
            "req_no":    r[idx["요청번호"]],
            "shipper":   r[idx["화주명"]],
            "pol":       r[idx["출발항"]],        # 출발
            "pod":       r[idx["도착항"]],        # 도착
            "cntr_type": r[idx["컨테이너타입"]],  # 타입
            "qty":       r[idx["수량"]],          # 수량
            "etd":       r[idx["희망출항일"]],
        })
    return requests


# ── 단계 2: 운임표에서 같은 구간+타입 단가 찾기 ────────────────────
def step2_load_rates(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    header = rows[0]
    idx = {name: i for i, name in enumerate(header)}
    rates = {}
    for r in rows[1:]:
        if r[idx["출발항"]] is None:
            continue
        key = (r[idx["출발항"]], r[idx["도착항"]], r[idx["컨테이너타입"]])
        rates[key] = {
            "unit_price":  r[idx["단가USD"]],
            "currency":    r[idx["통화"]],
            "valid_until": r[idx["유효기간"]],
        }
    return rates


def step2_find_rate(rates, req):
    key = (req["pol"], req["pod"], req["cntr_type"])  # 같은 구간 + 타입
    return rates.get(key)


# ── 단계 3: 합계 = 단가 × 수량 ─────────────────────────────────────
def step3_calc_total(unit_price, qty):
    return unit_price * qty


# ── 단계 4: 양식 빈칸 채워 건별 견적서로 저장 ──────────────────────
def _fmt(v):
    return v.strftime("%Y-%m-%d") if hasattr(v, "strftime") else str(v)


def _replace_in_doc(doc, mapping):
    def fix_para(p):
        full = "".join(run.text for run in p.runs)
        if "{{" not in full:
            return
        for k, v in mapping.items():
            full = full.replace(k, v)
        for i, run in enumerate(p.runs):
            run.text = full if i == 0 else ""
    for p in doc.paragraphs:
        fix_para(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix_para(p)


def step4_save_quote(template_path, out_dir, req, rate, total):
    os.makedirs(out_dir, exist_ok=True)
    doc = Document(template_path)
    mapping = {
        "{{REQ_NO}}":      str(req["req_no"]),
        "{{SHIPPER}}":     str(req["shipper"]),
        "{{QUOTE_DATE}}":  date.today().strftime("%Y-%m-%d"),
        "{{VALID_UNTIL}}": _fmt(rate["valid_until"]),
        "{{POL}}":         str(req["pol"]),
        "{{POD}}":         str(req["pod"]),
        "{{CNTR_TYPE}}":   str(req["cntr_type"]),
        "{{QTY}}":         str(req["qty"]),
        "{{ETD}}":         _fmt(req["etd"]),
        "{{UNIT_PRICE}}":  f'{rate["unit_price"]:,}',
        "{{TOTAL}}":       f"{total:,}",
    }
    _replace_in_doc(doc, mapping)
    out_path = os.path.join(out_dir, f'견적서_{req["req_no"]}.docx')
    doc.save(out_path)
    return out_path


# ── 메인: 항상 1→2→3→4 순서로만 실행 ──────────────────────────────
def main():
    # 1
    requests = step1_read_requests(SHIPPER_REQUEST)
    # 2 (운임표 로드)
    rates = step2_load_rates(RATE_TABLE)

    for req in requests:
        # 2 (해당 건 단가 조회)
        rate = step2_find_rate(rates, req)
        if rate is None:
            print(f'[건너뜀] {req["req_no"]}: 운임표에 해당 구간+타입 없음 '
                  f'({req["pol"]}->{req["pod"]} {req["cntr_type"]})')
            continue
        # 3
        total = step3_calc_total(rate["unit_price"], req["qty"])
        # 4
        out = step4_save_quote(QUOTE_TEMPLATE, OUTPUT_DIR, req, rate, total)
        print(f'[완료] {req["req_no"]}: 단가 {rate["unit_price"]:,} x 수량 {req["qty"]} '
              f'= 합계 {total:,} USD  ->  {os.path.basename(out)}')


if __name__ == "__main__":
    main()
